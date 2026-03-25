"""
generate_docx.py  –  HD Hauling & Grading Word proposal generator
Mirrors generate_proposal.py (PDF) as closely as python-docx allows.
Uses python-docx (pip install python-docx)
"""
import io, os, base64, tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

# ── Brand colours ───────────────────────────────────────────────────────────
RED      = RGBColor(0xCC, 0x00, 0x00)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BLACK    = RGBColor(0x11, 0x11, 0x11)
DARK     = RGBColor(0x1A, 0x1A, 0x1A)
LGRAY    = RGBColor(0xF4, 0xF4, 0xF4)
MGRAY    = RGBColor(0xCC, 0xCC, 0xCC)
DGRAY    = RGBColor(0x55, 0x55, 0x55)
COLHDR   = RGBColor(0x4A, 0x4A, 0x4A)
ROWALT   = RGBColor(0xEE, 0xEE, 0xEE)
TBLBORD  = RGBColor(0xCC, 0xCC, 0xCC)

_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_COVER = os.path.join(_DIR, 'hd_logo_cropped.png')
if not os.path.exists(LOGO_COVER):
    LOGO_COVER = os.path.join(_DIR, 'hd_logo.png')
LOGO_HEADER = os.path.join(_DIR, 'hd_logo.png')
if not os.path.exists(LOGO_HEADER):
    LOGO_HEADER = LOGO_COVER


# ── Helpers ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    """Set background shading on a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_width(cell, inches):
    """Force exact cell width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each arg is (size_eighth_pt, color_hex) or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            sz, clr = val
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), clr)
            el.set(qn('w:space'), '0')
            borders.append(el)
    tcPr.append(borders)


def _set_row_height(row, inches):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(inches * 1440)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def _set_table_borders(table, top=None, bottom=None, left=None, right=None,
                       insideH=None, insideV=None):
    """Set table-level borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for side, val in [('top', top), ('left', left), ('bottom', bottom),
                      ('right', right), ('insideH', insideH), ('insideV', insideV)]:
        if val is not None:
            sz, clr = val
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), clr)
            el.set(qn('w:space'), '0')
            borders.append(el)
        else:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
            borders.append(el)
    tblPr.append(borders)


def _merge_cells(table, row_idx, start_col, end_col):
    """Merge cells in a row from start_col to end_col (inclusive)."""
    row = table.rows[row_idx]
    row.cells[start_col].merge(row.cells[end_col])


def _add_run(para, text, bold=False, italic=False, size=10, color=None, font_name='Calibri'):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


def _set_para_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def _add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br_el = OxmlElement('w:br')
    br_el.set(qn('w:type'), 'page')
    run._r.append(br_el)
    return p


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in twips (1/1440 inch)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_page_number_footer(section):
    """Add 'Page X of Y' to the footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run('Page ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = 'Calibri'

    # PAGE field
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    run1 = p.add_run()
    run1._r.append(fld_char1)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run2.font.name = 'Calibri'
    run2._r.append(instr)
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run3 = p.add_run()
    run3._r.append(fld_char2)

    run4 = p.add_run(' of ')
    run4.font.size = Pt(8)
    run4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run4.font.name = 'Calibri'

    # NUMPAGES field
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'begin')
    run5 = p.add_run()
    run5._r.append(fld_char3)
    instr2 = OxmlElement('w:instrText')
    instr2.set(qn('xml:space'), 'preserve')
    instr2.text = ' NUMPAGES '
    run6 = p.add_run()
    run6.font.size = Pt(8)
    run6.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run6.font.name = 'Calibri'
    run6._r.append(instr2)
    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')
    run7 = p.add_run()
    run7._r.append(fld_char4)


def _add_header_to_section(section, date_str):
    """Add logo + 'PROPOSAL & CONTRACT' + date + black rule to header."""
    header = section.header
    header.is_linked_to_previous = False

    # Use a table for left-logo / right-text layout
    tbl = header.add_table(rows=1, cols=2, width=Inches(7.5))
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(5.5)
    _remove_table_borders(tbl)

    # Logo cell
    cell_logo = tbl.rows[0].cells[0]
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_HEADER):
        run = p_logo.add_run()
        run.add_picture(LOGO_HEADER, width=Inches(1.25))

    # Right cell — title and date
    cell_right = tbl.rows[0].cells[1]
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p_title = cell_right.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p_title, 'PROPOSAL & CONTRACT', bold=True, size=16, color=BLACK)

    p_date = cell_right.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(2)
    p_date.paragraph_format.space_after = Pt(4)
    _add_run(p_date, date_str, size=9, color=DGRAY)

    # Black horizontal rule below header
    p_rule = header.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after = Pt(0)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '8')
    bottom_border.set(qn('w:color'), '111111')
    bottom_border.set(qn('w:space'), '1')
    pBdr.append(bottom_border)
    pPr.append(pBdr)


# ── Cover Page ──────────────────────────────────────────────────────────────

def _build_cover_page(doc, data):
    """Build a cover page matching the PDF: centred logo, project name, subtitle, date,
    then prepared-by/prepared-for at the bottom."""

    # Big vertical spacer to push logo to upper-middle area
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Centred logo
    if os.path.exists(LOGO_COVER):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=0, after=6)
        run = p.add_run()
        run.add_picture(LOGO_COVER, width=Inches(4.0))

    # Project name — large bold centred
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=12, after=4)
    _add_run(p, data.get('project_name', 'Proposal'), bold=True, size=26, color=BLACK)

    # Thin grey divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=0, after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6')
    bottom_bdr.set(qn('w:color'), 'CCCCCC')
    bottom_bdr.set(qn('w:space'), '1')
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)
    # Set indent to narrow the line (simulate ~5.2" centered)
    pPr_ind = OxmlElement('w:ind')
    pPr_ind.set(qn('w:left'), '1440')   # 1 inch
    pPr_ind.set(qn('w:right'), '1440')
    pPr.append(pPr_ind)

    # Subtitle — "Proposal & Contract"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=10, after=4)
    _add_run(p, 'Proposal & Contract', size=18, color=DGRAY)

    # Date
    date_str = data.get('date', '')
    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=6, after=0)
        _add_run(p, date_str, size=13, color=RGBColor(0x99, 0x99, 0x99))

    # Push prepared by / prepared for to bottom of page
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Prepared by / Prepared for — two-column table
    # Count rows needed: label + name + company + email + phone
    max_rows = 5
    tbl = doc.add_table(rows=max_rows, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(3.5)
    tbl.columns[1].width = Inches(3.5)
    _remove_table_borders(tbl)

    # Row 0 — labels
    p = tbl.rows[0].cells[0].paragraphs[0]
    _add_run(p, 'Prepared by:', bold=True, size=10, color=BLACK)
    p = tbl.rows[0].cells[1].paragraphs[0]
    _add_run(p, 'Prepared for:', bold=True, size=10, color=BLACK)

    # Row 1 — names (bold)
    p = tbl.rows[1].cells[0].paragraphs[0]
    _set_para_spacing(p, before=4, after=1)
    _add_run(p, data.get('sender_name', ''), bold=True, size=10, color=DGRAY)
    p = tbl.rows[1].cells[1].paragraphs[0]
    _set_para_spacing(p, before=4, after=1)
    _add_run(p, data.get('client_name', ''), bold=True, size=10, color=DGRAY)

    # Row 2 — company
    p = tbl.rows[2].cells[0].paragraphs[0]
    _set_para_spacing(p, before=0, after=1)
    _add_run(p, data.get('company', 'HD Hauling & Grading'), size=9, color=DGRAY)
    p = tbl.rows[2].cells[1].paragraphs[0]
    _set_para_spacing(p, before=0, after=1)
    _add_run(p, data.get('client_company', ''), size=9, color=DGRAY)

    # Row 3 — email
    p = tbl.rows[3].cells[0].paragraphs[0]
    _set_para_spacing(p, before=0, after=1)
    _add_run(p, data.get('sender_email', ''), size=9, color=DGRAY)
    p = tbl.rows[3].cells[1].paragraphs[0]
    _set_para_spacing(p, before=0, after=1)
    _add_run(p, data.get('client_email', ''), size=9, color=DGRAY)

    # Row 4 — phone
    p = tbl.rows[4].cells[0].paragraphs[0]
    _set_para_spacing(p, before=0, after=0)
    _add_run(p, data.get('sender_phone', ''), size=9, color=DGRAY)
    p = tbl.rows[4].cells[1].paragraphs[0]
    _set_para_spacing(p, before=0, after=0)
    _add_run(p, data.get('client_phone', ''), size=9, color=DGRAY)


# ── Info Block ──────────────────────────────────────────────────────────────

def _build_info_block(doc, data):
    """3-column info block: project info | prepared by | prepared for.
    Matches the PDF info_block layout with vertical separator lines."""

    tbl = doc.add_table(rows=4, cols=3)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(3.15)
    tbl.columns[1].width = Inches(2.03)
    tbl.columns[2].width = Inches(2.32)

    # Top border on entire table, bottom border light gray
    _set_table_borders(tbl,
                       top=(8, '111111'),
                       bottom=(4, 'CCCCCC'),
                       left=None, right=None,
                       insideH=None, insideV=None)

    # ── Column 0: Project info ──
    # Row 0: project name
    cell = tbl.rows[0].cells[0]
    p = cell.paragraphs[0]
    _add_run(p, data.get('project_name', ''), bold=True, size=11, color=BLACK)
    _set_para_spacing(p, before=6, after=2)

    # Row 1: address
    cell = tbl.rows[1].cells[0]
    addr_parts = filter(None, [data.get('address', ''), data.get('city_state', '')])
    p = cell.paragraphs[0]
    _add_run(p, ', '.join(addr_parts), size=8, color=DGRAY)
    _set_para_spacing(p, before=0, after=2)

    # Row 2: spacer
    cell = tbl.rows[2].cells[0]
    p = cell.paragraphs[0]
    _set_para_spacing(p, before=0, after=0)

    # Row 3: date (red bold)
    cell = tbl.rows[3].cells[0]
    p = cell.paragraphs[0]
    _add_run(p, data.get('date', ''), bold=True, size=8, color=RED)
    _set_para_spacing(p, before=0, after=6)

    # ── Column 1: Prepared By ──
    cell = tbl.rows[0].cells[1]
    p = cell.paragraphs[0]
    _add_run(p, 'PREPARED BY', bold=True, size=7, color=RED)
    _set_para_spacing(p, before=6, after=2)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    cell = tbl.rows[1].cells[1]
    p = cell.paragraphs[0]
    _add_run(p, data.get('sender_name', ''), bold=True, size=9, color=BLACK)
    _set_para_spacing(p, before=0, after=1)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    cell = tbl.rows[2].cells[1]
    p = cell.paragraphs[0]
    _add_run(p, data.get('sender_email', ''), size=8, color=DGRAY)
    _set_para_spacing(p, before=0, after=1)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    cell = tbl.rows[3].cells[1]
    p = cell.paragraphs[0]
    _add_run(p, data.get('sender_phone', ''), size=8, color=DGRAY)
    _set_para_spacing(p, before=0, after=6)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    # ── Column 2: Prepared For ──
    cell = tbl.rows[0].cells[2]
    p = cell.paragraphs[0]
    _add_run(p, 'PREPARED FOR', bold=True, size=7, color=RED)
    _set_para_spacing(p, before=6, after=2)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    cell = tbl.rows[1].cells[2]
    p = cell.paragraphs[0]
    _add_run(p, data.get('client_name', ''), bold=True, size=9, color=BLACK)
    _set_para_spacing(p, before=0, after=1)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    cell = tbl.rows[2].cells[2]
    p = cell.paragraphs[0]
    _add_run(p, data.get('client_email', ''), size=8, color=DGRAY)
    _set_para_spacing(p, before=0, after=1)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    cell = tbl.rows[3].cells[2]
    p = cell.paragraphs[0]
    _add_run(p, data.get('client_phone', ''), size=8, color=DGRAY)
    _set_para_spacing(p, before=0, after=6)
    _set_cell_borders(cell, left=(4, 'CCCCCC'))

    # Add left padding to cols 1 & 2
    for row in tbl.rows:
        for ci in (1, 2):
            _set_cell_margins(row.cells[ci], left=200, top=0, bottom=0, right=86)


# ── Notes Block ─────────────────────────────────────────────────────────────

def _build_notes_block(doc, data):
    """Bordered box with 'Notes:' bold label followed by notes text."""
    notes = (data.get('notes', '') or '').strip()
    if not notes:
        return

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(7.5)
    _set_table_borders(tbl,
                       top=(4, 'CCCCCC'), bottom=(4, 'CCCCCC'),
                       left=(4, 'CCCCCC'), right=(4, 'CCCCCC'))

    cell = tbl.rows[0].cells[0]
    _set_cell_margins(cell, top=115, bottom=115, left=144, right=144)
    p = cell.paragraphs[0]
    _add_run(p, 'Notes:  ', bold=True, size=9, color=BLACK)
    _add_run(p, notes, size=9, color=DGRAY)
    _set_para_spacing(p, before=0, after=0, line=13)


# ── Bid Items Table ─────────────────────────────────────────────────────────

def _build_bid_table(doc, data):
    """Red banner header, dark column headers, alternating data rows."""
    line_items = data.get('line_items', [])
    num_data_rows = len(line_items)
    # rows: 0=red banner, 1=column headers, 2..N=data
    total_rows = 2 + num_data_rows
    tbl = doc.add_table(rows=total_rows, cols=5)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [3.75, 0.75, 0.75, 1.125, 1.125]
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = Inches(w)

    # Remove default borders, we'll set them manually
    _set_table_borders(tbl,
                       top=(4, 'CCCCCC'), bottom=(4, 'CCCCCC'),
                       left=(4, 'CCCCCC'), right=(4, 'CCCCCC'),
                       insideH=None, insideV=None)

    # ── Row 0: Red banner ──
    _merge_cells(tbl, 0, 0, 4)
    banner_cell = tbl.rows[0].cells[0]
    _set_cell_bg(banner_cell, 'CC0000')
    p = banner_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'BID ITEMS', bold=True, size=11, color=WHITE)
    _set_para_spacing(p, before=6, after=6)
    _set_cell_vertical_alignment(banner_cell, 'center')

    # ── Row 1: Column headers (dark gray background) ──
    headers = ['ITEM & DESCRIPTION', 'QTY', 'UNIT', 'PRICE', 'SUBTOTAL']
    aligns  = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
               WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
               WD_ALIGN_PARAGRAPH.RIGHT]
    for i, (h, align) in enumerate(zip(headers, aligns)):
        cell = tbl.rows[1].cells[i]
        _set_cell_bg(cell, '4A4A4A')
        p = cell.paragraphs[0]
        p.alignment = align
        _add_run(p, h, bold=True, size=8, color=WHITE)
        _set_para_spacing(p, before=5, after=5)
        _set_cell_vertical_alignment(cell, 'center')
    _set_row_height(tbl.rows[1], 0.28)

    # ── Data rows ──
    for idx, item in enumerate(line_items):
        row_idx = idx + 2
        row = tbl.rows[row_idx]

        # Alternating background
        bg = 'EEEEEE' if idx % 2 == 0 else 'FFFFFF'

        name  = item.get('name', '')
        desc  = item.get('description', '')
        qty   = item.get('qty', 0)
        unit  = item.get('unit', 'SY')
        price = item.get('price', 0)
        sub   = item.get('subtotal', 0)

        qty_s = f'{int(qty):,}' if isinstance(qty, (int, float)) and qty == int(qty) else str(qty)

        # Column 0: Item name (bold) + description (smaller gray)
        cell = row.cells[0]
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        _add_run(p, name, bold=True, size=9, color=BLACK)
        if desc:
            p2 = cell.add_paragraph()
            _add_run(p2, desc, size=8, color=RGBColor(0x77, 0x77, 0x77))
            _set_para_spacing(p2, before=1, after=0)
        _set_para_spacing(p, before=3, after=1 if desc else 3)

        # Columns 1-4
        vals = [qty_s, unit, f'${price:,.2f}', f'${sub:,.2f}']
        for ci, (v, align) in enumerate(zip(vals, aligns[1:]), start=1):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            _set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = align
            is_subtotal = (ci == 4)
            _add_run(p, v, bold=is_subtotal, size=9, color=BLACK)
            _set_para_spacing(p, before=3, after=3)

        # Light gray bottom border on data rows
        if idx < num_data_rows - 1:
            for ci in range(5):
                _set_cell_borders(row.cells[ci], bottom=(2, 'CCCCCC'))


# ── Contract Total Line ─────────────────────────────────────────────────────

def _build_total_line(doc, data):
    """Gray background bar with CONTRACT TOTAL on left, dollar amount on right, red bottom border."""
    total = data.get('total', 0)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(4.5)
    tbl.columns[1].width = Inches(3.0)

    _set_table_borders(tbl,
                       top=(4, 'CCCCCC'), bottom=(16, 'CC0000'),
                       left=None, right=None)

    row = tbl.rows[0]
    _set_row_height(row, 0.44)

    # Left cell
    cell = row.cells[0]
    _set_cell_bg(cell, 'F4F4F4')
    _set_cell_vertical_alignment(cell, 'center')
    _set_cell_margins(cell, left=144, right=0, top=0, bottom=0)
    p = cell.paragraphs[0]
    _add_run(p, 'CONTRACT TOTAL', bold=True, size=11, color=BLACK)

    # Right cell
    cell = row.cells[1]
    _set_cell_bg(cell, 'F4F4F4')
    _set_cell_vertical_alignment(cell, 'center')
    _set_cell_margins(cell, left=0, right=144, top=0, bottom=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p, f'${total:,.2f}', bold=True, size=11, color=BLACK)


# ── Site Plan Page ──────────────────────────────────────────────────────────

def _build_site_plan(doc, data):
    """Embed site plan image if provided, otherwise dashed placeholder box."""
    site_plan_b64 = data.get('site_plan_image', '')
    if not site_plan_b64:
        return

    _add_page_break(doc)

    if site_plan_b64 and ',' in site_plan_b64:
        try:
            img_data = base64.b64decode(site_plan_b64.split(',')[1])
            img_ext = 'png'
            if ';' in site_plan_b64:
                mime = site_plan_b64.split(';')[0]
                if '/' in mime:
                    img_ext = mime.split('/')[1].split('+')[0]
            with tempfile.NamedTemporaryFile(suffix=f'.{img_ext}', delete=False) as tmp:
                tmp.write(img_data)
                tmp_path = tmp.name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=12, after=0)
            run = p.add_run()
            run.add_picture(tmp_path, width=Inches(6.5))
            os.unlink(tmp_path)
            return
        except Exception:
            pass

    # Placeholder text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=200, after=6)
    _add_run(p, 'Site Plan / Drawing', bold=True, size=14, color=MGRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=0, after=4)
    _add_run(p, 'Replace this page with site plan image', size=10, color=MGRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=0, after=0)
    _add_run(p, 'or attach separately before sending to client', size=10, color=MGRAY)


# ── Terms & Conditions ──────────────────────────────────────────────────────

def _build_tc_section(doc, title, body_items):
    """One T&C section: red-left-border header on gray background, then body text/bullets."""
    # Header with red left border and gray background
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(7.5)
    _set_table_borders(tbl, top=None, bottom=(2, 'CCCCCC'), left=(32, 'CC0000'), right=None)

    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, 'F6F6F6')
    _set_cell_margins(cell, top=86, bottom=86, left=144, right=0)
    p = cell.paragraphs[0]
    _add_run(p, title, bold=True, size=10, color=RED)
    _set_para_spacing(p, before=0, after=0)

    # Body paragraphs
    for item in body_items:
        p = doc.add_paragraph()
        if item.startswith('\u2022') or item.startswith('•'):
            text = item.lstrip('\u2022•').strip()
            _set_para_spacing(p, before=1, after=4, line=14)
            p.paragraph_format.left_indent = Pt(16)
            _add_run(p, '- ' + text, size=8, color=DGRAY)
        else:
            _set_para_spacing(p, before=1, after=5, line=14)
            _add_run(p, item, size=8, color=DGRAY)

    # Small spacer after each section
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=3)


def _build_terms_and_conditions(doc):
    """All 20 T&C sections — content identical to generate_proposal.py tc_pages."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=6, after=6)
    _add_run(p, 'Terms & Conditions', bold=True, size=14, color=BLACK)

    # Red horizontal rule
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '8')
    bottom_bdr.set(qn('w:color'), 'CC0000')
    bottom_bdr.set(qn('w:space'), '1')
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)

    sections = [
        ('1. Contract Formation & Binding Agreement', [
            'This Proposal & Contract becomes legally binding upon execution by both the Customer/Purchaser and HD Hauling & Grading. Any conditions not expressly set forth herein shall not be recognized unless documented in writing and signed by authorized representatives of both parties. Verbal agreements, purchase orders, or prior understandings do not modify or supersede this contract unless incorporated by written amendment.',
        ]),
        ('2. Proposal Validity', [
            'Pricing in this proposal is valid for thirty (30) calendar days from the date of issuance. HD Hauling & Grading reserves the right to withdraw or modify this proposal if not executed within that period, including adjustments for material price changes.',
        ]),
        ('3. Scope of Work', [
            'HD Hauling & Grading\'s scope is limited to the paving, concrete, striping, and signage work explicitly described in the Bid Items section of this document. No additional work, modifications, or extensions of scope are included unless captured in a written, signed Change Order prior to commencement of that work.',
        ]),
        ('4. Change Orders', [
            'Any modification to the approved scope of work \u2014 including additions, deletions, substitutions, or design changes \u2014 requires a written Change Order executed by both parties before work begins. HD Hauling & Grading shall not be obligated to perform out-of-scope work without an approved Change Order and is not liable for delays caused by scope changes requested after contract execution.',
        ]),
        ('5. Site Access & Staging', [
            'The Customer shall provide HD Hauling & Grading with unobstructed vehicular access to the project site, a designated staging area for equipment and materials, and a safe haul route for loaded delivery trucks for the duration of work.',
            '\u2022 Delays, re-mobilizations, or standby time caused by restricted access, site conflicts with other trades, or unavailability of the work area will be billed at the applicable unit rates in the Paving Overage Unit Prices.',
            '\u2022 Customer is responsible for ensuring underground utilities are located and marked prior to the start of work. HD Hauling & Grading is not liable for damage to unmarked utilities.',
        ]),
        ('6. Subgrade Acceptance & Pavement Performance', [
            'HD Hauling & Grading is not responsible for pavement failure, cracking, settlement, or premature deterioration resulting from inadequate subgrade preparation, insufficient base compaction, poor drainage, or unsuitable sub-base materials outside HD Hauling & Grading\'s scope of work.',
            '\u2022 Prior to paving, the Customer or their designated representative is responsible for ensuring the subgrade and base course have been properly graded, compacted to NCDOT specifications, proof-rolled where required, and inspected. Commencement of paving constitutes Customer\'s acceptance of subgrade conditions.',
            '\u2022 Proof rolling, moisture content testing, and base course density testing are the responsibility of the Customer unless explicitly included in the scope of work.',
            '\u2022 If HD Hauling & Grading identifies conditions that may affect pavement performance, written notification will be provided. Customer\'s direction to proceed releases HD Hauling & Grading from performance liability related to those conditions.',
        ]),
        ('7. Materials & NCDOT Specifications', [
            'All asphalt materials furnished by HD Hauling & Grading shall conform to the applicable NCDOT Standard Specifications for Roads and Structures, current edition, unless otherwise specified in writing. Mix type, aggregate gradation, and binder content shall be per the mix design designated in the Bid Items.',
            '\u2022 Material substitutions required due to plant availability or supply chain disruptions will be communicated promptly. Functionally equivalent materials will be substituted at no additional cost.',
            '\u2022 HD Hauling & Grading does not guarantee long-term availability of specific mix designs or material sources.',
        ]),
        ('8. Weather & Temperature Conditions \u2014 Asphalt', [
            'Asphalt paving will not be performed under the following conditions: ambient or surface temperatures below 40\u00b0F and falling; during rain, sleet, or snow; when the base course contains standing water or frost; or when forecast conditions within four (4) hours are anticipated to compromise compaction or curing.',
            '\u2022 Schedule adjustments caused by weather are not grounds for price renegotiation, penalties, or liquidated damages against HD Hauling & Grading.',
        ]),
        ('9. Concrete Work Conditions', [
            'All concrete work shall be performed in accordance with applicable ACI standards and NCDOT specifications. Concrete will not be placed when ambient temperatures are below 40\u00b0F without approved cold-weather protection measures, or when temperatures exceed 90\u00b0F without appropriate hot-weather precautions.',
            '\u2022 Form layout, grade stakes, and joint locations must be approved by the Customer or their representative prior to placement. Once concrete is placed, corrections to layout or elevation are billable as additional work.',
            '\u2022 Cure time and form removal timing will be determined by HD Hauling & Grading based on ambient conditions and mix design requirements. Customer requests to accelerate form removal or trafficking of concrete prior to adequate cure are at Customer\'s sole risk.',
            '\u2022 HD Hauling & Grading is not responsible for surface defects, cracking, or scaling resulting from: improper curing practices by others, premature trafficking, freeze-thaw cycles, de-icing chemical application, or subgrade settlement outside HD Hauling & Grading\'s scope.',
        ]),
        ('10. Compaction & Quality', [
            'Asphalt pavement compaction shall meet NCDOT density requirements for the specified mix type. If compaction testing is required, the Customer is responsible for providing an independent testing agency.',
            '\u2022 HD Hauling & Grading is not liable for compaction failures resulting from: mix temperature loss during transport caused by factors outside its control, Customer-caused delays between delivery and laydown, or subgrade instability.',
        ]),
        ('11. Warranty', [
            'HD Hauling & Grading warrants all materials and workmanship for one (1) year from the date of substantial completion. This warranty covers defects in materials and workmanship performed directly by HD Hauling & Grading.',
            'This warranty expressly excludes:',
            '\u2022 Damage from petroleum products, chemical spills, or de-icing agents',
            '\u2022 Pavement failure from subgrade conditions not prepared by HD Hauling & Grading',
            '\u2022 Damage from vehicle loads exceeding pavement design capacity',
            '\u2022 Deterioration adjacent to a repaired area on maintenance projects',
            '\u2022 Pavement markings or signage not installed by HD Hauling & Grading',
            '\u2022 Normal wear, surface oxidation, and expected pavement aging',
            '\u2022 Damage from third parties, acts of God, or events beyond HD Hauling & Grading\'s control',
            'For maintenance and repair projects, the warranty applies only to the specific area(s) of new work.',
        ]),
        ('12. Traffic Control & Permits', [
            'If included in the Bid Items, HD Hauling & Grading will provide traffic control in general conformance with the MUTCD for the duration of active paving operations.',
            '\u2022 The Customer is responsible for all permits, right-of-way authorizations, and NCDOT lane closure approvals prior to the scheduled start of work.',
            '\u2022 ADA compliance determinations for pavement markings, curb ramps, and accessible routes are the responsibility of the Owner and Engineer of Record.',
        ]),
        ('13. Pavement Markings & Signage', [
            'Pavement markings will be installed per the approved plan or layout provided by the Customer. HD Hauling & Grading is not responsible for incorrect layouts resulting from inaccurate field dimensions or conflicting plans.',
            '\u2022 Thermoplastic markings require a minimum asphalt cure period before application.',
            '\u2022 Signage installation will follow locations and specifications provided. Sign content, ADA designation, and regulatory compliance are the Customer\'s responsibility.',
        ]),
        ('14. Limitation of Liability', [
            'HD Hauling & Grading\'s total liability under this contract, regardless of cause, shall not exceed the total contract value. In no event shall HD Hauling & Grading be liable for consequential, incidental, indirect, or punitive damages, including but not limited to loss of use, lost revenue, business interruption, or third-party claims arising from delays or defects.',
        ]),
        ('15. Payment Terms', [
            'All invoices are due Net 30 (thirty calendar days from the invoice date). Invoices will be submitted upon completion of each defined phase of work or on a monthly basis, whichever occurs first.',
            '\u2022 Balances not received within thirty (30) days accrue interest at 1.5% per month (18% annually).',
            '\u2022 Final payment is due within thirty (30) calendar days of the final completion invoice. Where applicable and agreed in writing, retention may be withheld per the terms of the prime contract, but shall be released no later than thirty (30) days after final completion and acceptance of HD Hauling & Grading\'s scope of work.',
            '\u2022 The individual executing this contract on behalf of the Customer/Purchaser provides a personal guarantee for full payment of all principal and accrued interest.',
        ]),
        ('16. Lien Rights', [
            'HD Hauling & Grading expressly reserves its right to file a Claim of Lien pursuant to N.C.G.S. Chapter 44A in the event of non-payment. Nothing herein constitutes a waiver of lien rights. In the event legal action is required, the Customer shall be responsible for all reasonable attorney\'s fees and collection costs per N.C.G.S. \u00a7 44A-35.',
        ]),
        ('17. Material Pricing & Availability', [
            'Due to volatility in liquid asphalt, aggregate, and fuel markets, material costs may be adjusted to reflect prevailing market rates if costs increase more than ten percent (10%) from the proposal date. Written notice will be provided prior to any adjustment.',
            '\u2022 HD Hauling & Grading is not liable for delays caused by plant shutdowns, material shortages, or supplier issues.',
        ]),
        ('18. Force Majeure', [
            'HD Hauling & Grading shall not be liable for delays or failure to perform caused by circumstances beyond its reasonable control, including acts of God, severe weather, labor disputes, government actions, supply chain disruptions, fuel shortages, or public health emergencies. The schedule will be extended by a reasonable period and pricing may be subject to renegotiation.',
        ]),
        ('19. Entire Agreement', [
            'This Proposal & Contract is the entire agreement between the parties with respect to the work described herein. It supersedes all prior proposals, representations, and understandings. No terms printed on the Customer\'s purchase orders shall apply unless explicitly incorporated by written amendment signed by both parties.',
        ]),
        ('20. Dispute Resolution', [
            'The parties agree to attempt to resolve any dispute arising under this Proposal & Contract through good-faith negotiation prior to initiating formal proceedings. If a dispute cannot be resolved through negotiation within thirty (30) days of written notice, either party may pursue resolution through binding arbitration administered under the rules of the American Arbitration Association, or by filing a claim in a court of competent jurisdiction in the State of North Carolina.',
            '\u2022 This Proposal & Contract shall be governed by and construed in accordance with the laws of the State of North Carolina, without regard to its conflict of law provisions.',
            '\u2022 The prevailing party in any arbitration or litigation arising from this Proposal & Contract shall be entitled to recover reasonable attorney\u2019s fees and costs from the non-prevailing party.',
        ]),
    ]

    for title, body in sections:
        _build_tc_section(doc, title, body)


# ── Approval Page ───────────────────────────────────────────────────────────

def _build_approval_page(doc, data):
    """Client approval & authorization page matching the PDF."""
    total = data.get('total', 0)

    # ── Red banner header ──
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(7.5)
    _set_table_borders(tbl, top=None, bottom=None, left=None, right=None)

    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, 'CC0000')
    _set_cell_margins(cell, top=100, bottom=100, left=115, right=0)
    p = cell.paragraphs[0]
    _add_run(p, 'Client Approval & Authorization', bold=True, size=10, color=WHITE)
    _set_para_spacing(p, before=0, after=0)

    # Spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=8)

    # ── Project summary box ──
    sum_tbl = doc.add_table(rows=2, cols=4)
    sum_tbl.autofit = False
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sum_tbl.columns[0].width = Inches(0.9)
    sum_tbl.columns[1].width = Inches(2.85)
    sum_tbl.columns[2].width = Inches(0.9)
    sum_tbl.columns[3].width = Inches(2.85)

    _set_table_borders(sum_tbl,
                       top=(4, 'CCCCCC'), bottom=(4, 'CCCCCC'),
                       left=(4, 'CCCCCC'), right=(4, 'CCCCCC'),
                       insideH=(2, 'CCCCCC'), insideV=None)

    proj   = data.get('project_name', '')
    client = data.get('client_name', '')
    addr   = ', '.join(filter(None, [data.get('address', ''), data.get('city_state', '')]))
    date   = data.get('date', '')

    sum_data = [
        [('PROJECT', proj), ('CLIENT', client)],
        [('ADDRESS', addr), ('DATE', date)],
    ]

    for ri, row_data in enumerate(sum_data):
        for ci, (label, value) in enumerate(row_data):
            lbl_cell = sum_tbl.rows[ri].cells[ci * 2]
            val_cell = sum_tbl.rows[ri].cells[ci * 2 + 1]

            _set_cell_bg(lbl_cell, 'F4F4F4')
            _set_cell_bg(val_cell, 'F4F4F4')
            _set_cell_margins(lbl_cell, top=86, bottom=86, left=115, right=0)
            _set_cell_margins(val_cell, top=86, bottom=86, left=58, right=115)
            _set_cell_vertical_alignment(lbl_cell, 'center')
            _set_cell_vertical_alignment(val_cell, 'center')

            p = lbl_cell.paragraphs[0]
            _add_run(p, label, bold=True, size=8, color=DGRAY)
            _set_para_spacing(p, before=0, after=0)

            p = val_cell.paragraphs[0]
            _add_run(p, value, size=9, color=BLACK)
            _set_para_spacing(p, before=0, after=0)

            # Vertical separator before the second pair
            if ci == 1:
                _set_cell_borders(lbl_cell, left=(2, 'CCCCCC'))

    # Spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=8)

    # ── Approved contract value bar (matches CONTRACT TOTAL style) ──
    amt_tbl = doc.add_table(rows=1, cols=2)
    amt_tbl.autofit = False
    amt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    amt_tbl.columns[0].width = Inches(4.5)
    amt_tbl.columns[1].width = Inches(3.0)

    _set_table_borders(amt_tbl,
                       top=(4, 'CCCCCC'), bottom=(16, 'CC0000'),
                       left=None, right=None)

    row = amt_tbl.rows[0]
    _set_row_height(row, 0.44)

    cell = row.cells[0]
    _set_cell_bg(cell, 'F4F4F4')
    _set_cell_vertical_alignment(cell, 'center')
    _set_cell_margins(cell, left=144, right=0, top=0, bottom=0)
    p = cell.paragraphs[0]
    _add_run(p, 'APPROVED CONTRACT VALUE', bold=True, size=11, color=BLACK)

    cell = row.cells[1]
    _set_cell_bg(cell, 'F4F4F4')
    _set_cell_vertical_alignment(cell, 'center')
    _set_cell_margins(cell, left=0, right=144, top=0, bottom=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p, f'${total:,.2f}', bold=True, size=11, color=BLACK)

    # Spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=10)

    # ── Authorization language (italic, centered) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=0, after=10, line=13)
    _add_run(p,
        'By signing below, Client authorizes HD Hauling & Grading to proceed with the work '
        'described in this Proposal & Contract and agrees to be bound by all terms and '
        'conditions set forth herein.',
        italic=True, size=8, color=DGRAY)

    # Spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=10)

    # ── Bilateral signature block ──
    sig_tbl = doc.add_table(rows=5, cols=2)
    sig_tbl.autofit = False
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.columns[0].width = Inches(3.75)
    sig_tbl.columns[1].width = Inches(3.75)

    _set_table_borders(sig_tbl,
                       top=(4, 'CCCCCC'), bottom=(4, 'CCCCCC'),
                       left=None, right=None,
                       insideH=None, insideV=None)

    sig_data = [
        ('HD Hauling & Grading', 'Client / Authorized Representative'),
        ('Authorized Signature: ___________________________', 'Authorized Signature: ___________________________'),
        ('Printed Name: _________________________________', 'Printed Name: _________________________________'),
        ('Title: _________________________________________', 'Title: _________________________________________'),
        ('Date: __________________________________________', 'Date: __________________________________________'),
    ]

    for ri, (left_text, right_text) in enumerate(sig_data):
        for ci, text in enumerate([left_text, right_text]):
            cell = sig_tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            is_header = (ri == 0)
            _add_run(p, text, bold=is_header, size=9, color=BLACK)
            _set_para_spacing(p, before=5, after=5)
            _set_cell_margins(cell, left=58, right=58, top=0, bottom=0)

    # Spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=12)

    # ── Additional Unit Prices table (if any) ──
    unit_items = data.get('unit_prices', [])
    if unit_items:
        num_rows = 2 + len(unit_items)
        up_tbl = doc.add_table(rows=num_rows, cols=2)
        up_tbl.autofit = False
        up_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        up_tbl.columns[0].width = Inches(5.85)
        up_tbl.columns[1].width = Inches(1.65)

        _set_table_borders(up_tbl,
                           top=(4, 'CCCCCC'), bottom=None,
                           left=(4, 'CCCCCC'), right=(4, 'CCCCCC'),
                           insideH=None, insideV=None)

        # Row 0: Red banner
        _merge_cells(up_tbl, 0, 0, 1)
        banner_cell = up_tbl.rows[0].cells[0]
        _set_cell_bg(banner_cell, 'CC0000')
        p = banner_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, 'Additional Unit Prices', bold=True, size=10, color=WHITE)
        _set_para_spacing(p, before=6, after=6)

        # Row 1: Column headers
        hdr_texts = ['Description', 'Unit Rate']
        hdr_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
        for ci, (h, al) in enumerate(zip(hdr_texts, hdr_aligns)):
            cell = up_tbl.rows[1].cells[ci]
            _set_cell_bg(cell, '4A4A4A')
            _set_cell_vertical_alignment(cell, 'center')
            p = cell.paragraphs[0]
            p.alignment = al
            _add_run(p, h, bold=True, size=8, color=WHITE)
            _set_para_spacing(p, before=5, after=5)
        _set_row_height(up_tbl.rows[1], 0.28)

        # Data rows
        for idx, item in enumerate(unit_items):
            ri = idx + 2
            row = up_tbl.rows[ri]
            bg = 'EEEEEE' if idx % 2 == 0 else 'FFFFFF'

            cell = row.cells[0]
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell, left=115, right=0, top=0, bottom=0)
            p = cell.paragraphs[0]
            _add_run(p, item.get('name', ''), size=9, color=BLACK)
            _set_para_spacing(p, before=4, after=4)

            cell = row.cells[1]
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell, left=0, right=115, top=0, bottom=0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_run(p, f'${item.get("rate", 0):,.2f}', bold=True, size=9, color=BLACK)
            _set_para_spacing(p, before=4, after=4)

            # Border between data rows
            if idx < len(unit_items) - 1:
                for ci in range(2):
                    _set_cell_borders(row.cells[ci], bottom=(2, 'CCCCCC'))

        # Red bottom border on last row
        last_row = up_tbl.rows[num_rows - 1]
        for ci in range(2):
            _set_cell_borders(last_row.cells[ci], bottom=(12, 'CC0000'))


# ── Main build function ─────────────────────────────────────────────────────

def build(data, out_path=None):
    """Build a Word proposal document matching the PDF layout.

    Args:
        data: dict with proposal data (same payload as generate_proposal.py)
        out_path: if provided, save to this file path; otherwise return BytesIO buffer
    """
    doc = Document()

    # ── Page setup: Letter, 0.5" L/R, ~1" top, 0.6" bottom ──
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(0.5)
    section.right_margin  = Inches(0.5)
    section.top_margin    = Inches(0.6)   # Cover page — no header needed
    section.bottom_margin = Inches(0.6)

    # No header on cover page, but add footer page numbers
    section.different_first_page_header_footer = False
    _add_page_number_footer(section)

    date_str = data.get('date', '')

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 1 — COVER
    # ════════════════════════════════════════════════════════════════════════
    _build_cover_page(doc, data)

    # ── New section for pages 2+ (with header) ──
    new_section = doc.add_section()
    new_section.page_width  = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.left_margin   = Inches(0.5)
    new_section.right_margin  = Inches(0.5)
    new_section.top_margin    = Inches(1.05)
    new_section.bottom_margin = Inches(0.6)

    _add_header_to_section(new_section, date_str)
    _add_page_number_footer(new_section)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 2+ — INFO BLOCK, NOTES, BID TABLE, TOTAL
    # ════════════════════════════════════════════════════════════════════════

    _build_info_block(doc, data)

    # Small spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=6)

    _build_notes_block(doc, data)

    # Small spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=6)

    _build_bid_table(doc, data)

    # Small spacer
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=4)

    _build_total_line(doc, data)

    # ════════════════════════════════════════════════════════════════════════
    # SITE PLAN (optional)
    # ════════════════════════════════════════════════════════════════════════
    _build_site_plan(doc, data)

    # ════════════════════════════════════════════════════════════════════════
    # TERMS & CONDITIONS
    # ════════════════════════════════════════════════════════════════════════
    _add_page_break(doc)
    _build_terms_and_conditions(doc)

    # ════════════════════════════════════════════════════════════════════════
    # APPROVAL PAGE
    # ════════════════════════════════════════════════════════════════════════
    _add_page_break(doc)
    _build_approval_page(doc, data)

    # ── Save ──
    if out_path:
        doc.save(out_path)
    else:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


if __name__ == '__main__':
    import json, sys
    data = json.loads(sys.argv[1]) if len(sys.argv) > 1 else {}
    out = sys.argv[2] if len(sys.argv) > 2 else 'test_proposal.docx'
    build(data, out)
    print(f'Written: {out}')
